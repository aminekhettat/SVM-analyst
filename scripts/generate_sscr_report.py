"""Generate a scientific DOCX report on Single Shunt Current Reconstruction (SSCR).

IEEE Transaction journal paper style.
Run from the project root:
    python scripts/generate_sscr_report.py

Output:
    docs/SSCR_Technical_Report.docx
"""

from __future__ import annotations

import io
import math
import os
import sys

import matplotlib

matplotlib.use("Agg")
import matplotlib.patches as mpatches
import matplotlib.pyplot as plt
import matplotlib.ticker as ticker
import numpy as np
from matplotlib.lines import Line2D
from matplotlib.patches import FancyArrowPatch, Rectangle

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _fig_to_stream(fig: plt.Figure, dpi: int = 200) -> io.BytesIO:
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=dpi, bbox_inches="tight")
    buf.seek(0)
    plt.close(fig)
    return buf


def _heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def _bold(para, text: str):
    run = para.add_run(text)
    run.bold = True
    return run


def _italic(para, text: str):
    run = para.add_run(text)
    run.italic = True
    return run


def _add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)


def _set_col_widths(table, widths_inches):
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths_inches):
                cell.width = Inches(widths_inches[idx])


def _shade_cell(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


# ---------------------------------------------------------------------------
# Figure generators
# ---------------------------------------------------------------------------

PHASE_COLORS = {"A": "#1f77b4", "B": "#ff7f0e", "C": "#2ca02c"}

SECTOR_COLORS = [
    "#ffd6d6",  # S1 red-ish
    "#ffecd6",  # S2 orange-ish
    "#fffbd6",  # S3 yellow-ish
    "#d6ffda",  # S4 green-ish
    "#d6f0ff",  # S5 blue-ish
    "#eld6ff",  # S6 purple-ish
]
SECTOR_COLORS = [
    "#ffd6d6",
    "#ffecd6",
    "#ffffd6",
    "#d6ffd6",
    "#d6f0ff",
    "#edd6ff",
]


def fig_inverter_topology() -> plt.Figure:
    """Three-phase half-bridge VSI with single DC-bus shunt."""
    fig, ax = plt.subplots(figsize=(7.5, 4.5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 8)
    ax.axis("off")
    ax.set_facecolor("white")

    # DC bus rails
    ax.plot([0.5, 9.5], [7, 7], "k-", lw=2)  # +VDC rail
    ax.plot([0.5, 9.5], [1, 1], "k-", lw=2)  # -VDC rail

    # DC bus labels
    ax.text(
        0.2,
        7,
        "+VDC",
        ha="right",
        va="center",
        fontsize=10,
        fontweight="bold",
        color="#c00000",
    )
    ax.text(
        0.2,
        1,
        "−VDC / GND",
        ha="right",
        va="center",
        fontsize=10,
        fontweight="bold",
        color="#c00000",
    )

    # Single shunt resistor on lower DC rail
    sx, sy = 1.5, 1.0
    ax.plot([sx - 0.3, sx - 0.3], [0.3, sy], "k-", lw=2)
    ax.plot([sx + 0.3, sx + 0.3], [0.3, sy], "k-", lw=2)
    # Zigzag resistor symbol
    xs = np.linspace(sx - 0.3, sx + 0.3, 8)
    ys = np.array([0.3, 0.45, 0.3, 0.45, 0.3, 0.45, 0.3, 0.45])
    ax.plot(xs, ys, "k-", lw=2)
    ax.text(sx, 0.1, r"$R_{shunt}$", ha="center", fontsize=9, color="#8B0000")

    # Battery / DC source on left
    ax.plot([0.5, 0.5], [1, 7], "k-", lw=1.5)
    ax.add_patch(plt.Rectangle((0.15, 3.8), 0.7, 0.4, color="gray"))
    ax.add_patch(plt.Rectangle((0.25, 3.6), 0.5, 0.2, color="gray"))
    ax.text(0.5, 3.1, "VDC", ha="center", fontsize=9, fontweight="bold")

    # Three half-bridges
    phases = ["A", "B", "C"]
    bridge_x = [3.5, 5.5, 7.5]
    phase_colors_hex = ["#1f77b4", "#ff7f0e", "#2ca02c"]

    for i, (ph, bx, col) in enumerate(zip(phases, bridge_x, phase_colors_hex)):
        # Vertical bus bar
        ax.plot([bx, bx], [1, 7], color="#aaa", lw=1, ls="--")

        # High-side switch (IGBT Q_xH)
        sw_y_h = 5.5
        ax.add_patch(
            plt.Polygon(
                [
                    [bx - 0.35, sw_y_h + 0.5],
                    [bx + 0.35, sw_y_h + 0.5],
                    [bx, sw_y_h - 0.3],
                ],
                color=col,
                alpha=0.7,
            )
        )
        ax.plot([bx, bx], [7, sw_y_h + 0.5], color=col, lw=2)
        ax.plot([bx, bx], [sw_y_h - 0.3, 4.8], color=col, lw=2)
        ax.text(bx + 0.5, sw_y_h + 0.1, f"$Q_{{{ph}H}}$", fontsize=9, color=col)

        # Freewheeling diode high side
        ax.annotate(
            "",
            xy=(bx - 0.55, sw_y_h + 0.4),
            xytext=(bx - 0.55, sw_y_h - 0.2),
            arrowprops=dict(arrowstyle="-|>", color=col, lw=1.5),
        )

        # Low-side switch (IGBT Q_xL)
        sw_y_l = 2.5
        ax.add_patch(
            plt.Polygon(
                [
                    [bx - 0.35, sw_y_l - 0.5],
                    [bx + 0.35, sw_y_l - 0.5],
                    [bx, sw_y_l + 0.3],
                ],
                color=col,
                alpha=0.7,
            )
        )
        ax.plot([bx, bx], [1, sw_y_l - 0.5], color=col, lw=2)
        ax.plot([bx, bx], [sw_y_l + 0.3, 4.0], color=col, lw=2)
        ax.text(bx + 0.5, sw_y_l - 0.1, f"$Q_{{{ph}L}}$", fontsize=9, color=col)

        # Freewheeling diode low side
        ax.annotate(
            "",
            xy=(bx - 0.55, sw_y_l - 0.4),
            xytext=(bx - 0.55, sw_y_l + 0.2),
            arrowprops=dict(arrowstyle="-|>", color=col, lw=1.5),
        )

        # Phase output mid-point
        ax.plot([bx, bx + 1.2], [4.4, 4.4], color=col, lw=2)
        # Motor phase label
        ax.text(
            bx + 1.4,
            4.4,
            f"$i_{ph}$",
            ha="left",
            va="center",
            fontsize=12,
            color=col,
            fontweight="bold",
        )
        ax.text(
            bx,
            8.0 if i == 0 else (8.0 if i == 1 else 8.0),
            f"Leg {ph}",
            ha="center",
            fontsize=10,
            color=col,
            fontweight="bold",
        )

    # Motor (simplified as 3-phase star winding)
    mx, my = 9.0, 4.4
    ax.add_patch(plt.Circle((mx, my), 0.5, fill=False, ec="gray", lw=2))
    ax.text(
        mx,
        my,
        "M",
        ha="center",
        va="center",
        fontsize=11,
        color="gray",
        fontweight="bold",
    )

    # Shunt voltage label
    ax.annotate(
        "",
        xy=(1.8, 0.6),
        xytext=(1.8, 1.0),
        arrowprops=dict(arrowstyle="<->", color="#8B0000", lw=1.5),
    )
    ax.text(
        2.1,
        0.8,
        r"$v_{shunt} = R_{shunt} \cdot i_{dc}$",
        fontsize=8,
        color="#8B0000",
        va="center",
    )

    ax.set_title(
        "Fig. 1 – Three-Phase Half-Bridge Voltage Source Inverter with Single DC-Bus Shunt Resistor",
        fontsize=10,
        pad=12,
    )
    fig.tight_layout()
    return fig


def fig_center_aligned_pwm() -> plt.Figure:
    """Detailed center-aligned PWM timing with W1/W2 acquisition windows."""
    fig, axes = plt.subplots(4, 1, figsize=(8.5, 7), sharex=True)
    fig.subplots_adjust(hspace=0.08, left=0.12, right=0.97, top=0.92, bottom=0.09)

    T = 1.0  # normalized period
    t = np.linspace(0, T, 2000)

    # Carrier
    carrier = 1.0 - 2.0 * np.abs(2.0 * (t / T) - 1.0)

    # Duty cycles for this illustration (sector 1: Da > Db > Dc)
    Da, Db, Dc = 0.75, 0.55, 0.30

    def center_pulse(D, t_arr, T_):
        # Center aligned: ON from T/2*(1-D) to T/2*(1+D)
        on_start = T_ / 2.0 * (1.0 - D)
        on_end = T_ / 2.0 * (1.0 + D)
        return np.where((t_arr >= on_start) & (t_arr <= on_end), 1.0, 0.0)

    pa = center_pulse(Da, t, T)
    pb = center_pulse(Db, t, T)
    pc = center_pulse(Dc, t, T)

    # Annotated times
    ton_a = T / 2 * (1 - Da)
    ton_b = T / 2 * (1 - Db)
    ton_c = T / 2 * (1 - Dc)
    toff_a = T / 2 * (1 + Da)
    toff_b = T / 2 * (1 + Db)
    toff_c = T / 2 * (1 + Dc)

    colors = [PHASE_COLORS["A"], PHASE_COLORS["B"], PHASE_COLORS["C"]]
    labels = ["PWM A", "PWM B", "PWM C"]
    pulses = [pa, pb, pc]
    duties = [Da, Db, Dc]

    for ax, col, lbl, pulse, D in zip(axes[:3], colors, labels, pulses, duties):
        ax.fill_between(t, 0, pulse, alpha=0.35, color=col)
        ax.plot(t, pulse, color=col, lw=1.5)
        ax.set_ylim(-0.15, 1.35)
        ax.set_yticks([0, 1])
        ax.set_yticklabels(["0", "1"], fontsize=8)
        ax.set_ylabel(lbl, fontsize=9, rotation=0, labelpad=42)
        ax.grid(True, alpha=0.3)
        ax.axhline(0, color="k", lw=0.5)

    # Shunt current on bottom axis
    ax_sh = axes[3]

    # Define shunt current state:
    # [0, ton_c] → 0 (all off or freewheeling, depends on current direction)
    # [ton_c, ton_b] → W2 window: Ia+Ib=−Ic visible
    # [ton_b, ton_a] → W1 only Ia window
    # [ton_a, T/2] → all ON (zero vector region)

    W1_start = ton_b  # first half
    W1_end = ton_a

    W2_start = ton_c
    W2_end = ton_b

    # shade W1 (green)
    for ax in axes[:3]:
        ax.axvspan(W1_start, W1_end, alpha=0.18, color="#00aa44", zorder=0)
        ax.axvspan(W2_start, W2_end, alpha=0.18, color="#aa4400", zorder=0)
        # Mirror in second half
        ax.axvspan(T - W1_end, T - W1_start, alpha=0.18, color="#00aa44", zorder=0)
        ax.axvspan(T - W2_end, T - W2_start, alpha=0.18, color="#aa4400", zorder=0)

    # Shunt current waveform
    i_sh = np.zeros_like(t)
    for k, tk in enumerate(t):
        # First half of period
        if W2_start <= tk < W2_end:
            i_sh[k] = -0.6  # represents −Ic
        elif W1_start <= tk < W1_end:
            i_sh[k] = 0.9  # represents Ia
        elif T / 2 - (W1_end - W1_start) <= tk < T / 2:
            i_sh[k] = 0.0  # all ON, zero state
        # Mirror second half
        elif (T - W1_end) <= tk < (T - W1_start):
            i_sh[k] = 0.9
        elif (T - W2_end) <= tk < (T - W2_start):
            i_sh[k] = -0.6
        else:
            i_sh[k] = 0.0

    ax_sh.plot(t, i_sh, color="#9467bd", lw=2)
    ax_sh.axhline(0, color="k", lw=0.5)
    ax_sh.axhspan(W1_start, W1_end, alpha=0.18, color="#00aa44")
    ax_sh.axvspan(W1_start, W1_end, alpha=0.18, color="#00aa44", zorder=0)
    ax_sh.axvspan(W2_start, W2_end, alpha=0.18, color="#aa4400", zorder=0)
    ax_sh.axvspan(T - W1_end, T - W1_start, alpha=0.18, color="#00aa44", zorder=0)
    ax_sh.axvspan(T - W2_end, T - W2_start, alpha=0.18, color="#aa4400", zorder=0)
    ax_sh.set_ylim(-1.4, 1.4)
    ax_sh.set_yticks([-0.6, 0, 0.9])
    ax_sh.set_yticklabels([r"$-I_C$", "0", r"$I_A$"], fontsize=8)
    ax_sh.set_ylabel(r"$i_{shunt}$", fontsize=9, rotation=0, labelpad=42)
    ax_sh.grid(True, alpha=0.3)

    # Annotate windows
    ax0 = axes[0]
    mid_w1 = (W1_start + W1_end) / 2
    mid_w2 = (W2_start + W2_end) / 2
    ax0.annotate(
        r"$W_1$",
        xy=(mid_w1, 1.2),
        ha="center",
        fontsize=9,
        color="#005500",
        fontweight="bold",
    )
    ax0.annotate(
        r"$W_2$",
        xy=(mid_w2, 1.2),
        ha="center",
        fontsize=9,
        color="#884400",
        fontweight="bold",
    )

    # T/2 line
    for ax in axes:
        ax.axvline(T / 2, color="k", lw=1, ls="--", alpha=0.5)

    axes[-1].set_xlabel(r"Time ($t / T_{PWM}$)", fontsize=10)
    axes[-1].set_xlim(0, T)
    axes[-1].set_xticks(
        [0, ton_c, ton_b, ton_a, 0.5, T - ton_a, T - ton_b, T - ton_c, T]
    )
    axes[-1].set_xticklabels(
        [
            "0",
            r"$t_{on,C}$",
            r"$t_{on,B}$",
            r"$t_{on,A}$",
            r"$T/2$",
            r"$t_{off,A}$",
            r"$t_{off,B}$",
            r"$t_{off,C}$",
            r"$T$",
        ],
        fontsize=7.5,
        rotation=45,
    )

    patch_w1 = mpatches.Patch(color="#00aa44", alpha=0.4, label=r"$W_1$ — sample $I_A$")
    patch_w2 = mpatches.Patch(
        color="#aa4400", alpha=0.4, label=r"$W_2$ — sample $-I_C$"
    )
    fig.legend(handles=[patch_w1, patch_w2], loc="upper right", fontsize=9)

    fig.suptitle(
        "Fig. 2 – Center-Aligned PWM Timing (Sector 1: $D_A > D_B > D_C$)\n"
        r"Acquisition windows $W_1$ and $W_2$ highlighted",
        fontsize=10,
    )
    return fig


def fig_window_geometry() -> plt.Figure:
    """Half-period zoom showing W1, W2, dead-time, and minimum observable threshold."""
    fig, ax = plt.subplots(figsize=(8.5, 3.5))
    T_half = 50.0  # µs
    dead = 2.0  # µs
    adc_min = 1.5  # µs

    Da, Db, Dc = 0.75, 0.55, 0.30

    ton_a = T_half * (1 - Da)
    ton_b = T_half * (1 - Db)
    ton_c = T_half * (1 - Dc)

    # Timeline bar
    ax.axhline(0.5, color="k", lw=1.5, xmin=ton_c / T_half, xmax=1)

    # Phase labels at switch-on times
    for t_on, lbl, col in zip(
        [ton_a, ton_b, ton_c],
        [r"$t_{on,A}$", r"$t_{on,B}$", r"$t_{on,C}$"],
        [PHASE_COLORS["A"], PHASE_COLORS["B"], PHASE_COLORS["C"]],
    ):
        ax.axvline(t_on, 0.3, 0.7, color=col, lw=1.5)
        ax.text(t_on, 0.75, lbl, ha="center", fontsize=9, color=col)

    # W1 region (with dead time eaten)
    w1_ideal = ton_b - ton_a
    w1_eff = w1_ideal - dead
    ax.add_patch(
        Rectangle(
            (ton_a, 0.1),
            dead,
            0.3,
            color="#e04040",
            alpha=0.7,
            label=r"Dead time $t_d$",
        )
    )
    ax.add_patch(
        Rectangle(
            (ton_a + dead, 0.1),
            w1_eff if w1_eff > 0 else 0,
            0.3,
            color="#00aa44",
            alpha=0.6,
            label=r"$W_1^{eff}$",
        )
    )
    mid_w1 = ton_a + dead + max(w1_eff, 0) / 2
    if w1_eff > 0:
        ax.text(
            mid_w1,
            0.47,
            rf"$W_1^{{eff}}={w1_eff:.1f}$ µs",
            ha="center",
            fontsize=8.5,
            color="#005500",
            fontweight="bold",
        )

    # W2 region
    w2_ideal = ton_a - ton_b  # ton_c to ton_b
    w2_ideal = ton_b - ton_c
    w2_eff = w2_ideal - dead
    ax.add_patch(Rectangle((ton_c, 0.1), dead, 0.3, color="#e04040", alpha=0.7))
    ax.add_patch(
        Rectangle(
            (ton_c + dead, 0.1),
            w2_eff if w2_eff > 0 else 0,
            0.3,
            color="#aa5500",
            alpha=0.5,
            label=r"$W_2^{eff}$",
        )
    )
    mid_w2 = ton_c + dead + max(w2_eff, 0) / 2
    if w2_eff > 0:
        ax.text(
            mid_w2,
            0.47,
            rf"$W_2^{{eff}}={w2_eff:.1f}$ µs",
            ha="center",
            fontsize=8.5,
            color="#883300",
            fontweight="bold",
        )

    # Minimum observable threshold line
    ax.axhline(0.25, color="purple", lw=1, ls=":", alpha=0.7)
    ax.annotate(
        r"$t_{acq,min}$",
        xy=(ton_c + dead + adc_min, 0.25),
        xytext=(ton_c + dead + adc_min + 3, 0.28),
        fontsize=8,
        color="purple",
        arrowprops=dict(arrowstyle="->", color="purple", lw=1),
    )

    ax.set_xlim(0, T_half)
    ax.set_ylim(0, 1)
    ax.set_xlabel(r"Time within half-period ($\mu$s)", fontsize=10)
    ax.set_yticks([])
    ax.set_title(
        r"Fig. 3 – Half-Period Window Analysis: Dead Time Consumption and Effective Acquisition Windows "
        "\n"
        r"($T_{half}=50\,\mu s$, $t_d=2\,\mu s$, $D_A=0.75$, $D_B=0.55$, $D_C=0.30$)",
        fontsize=9.5,
    )
    ax.legend(loc="upper right", fontsize=9)
    ax.grid(True, alpha=0.3)
    fig.tight_layout()
    return fig


def fig_svm_sectors() -> plt.Figure:
    """SVM hexagon with 6 sectors coloured and duty-cycle ordering given."""
    fig, ax = plt.subplots(figsize=(6.0, 5.8))
    ax.set_aspect("equal")
    ax.set_facecolor("#f8f8f8")

    radius = 1.0
    angles_v = np.linspace(0, 2 * np.pi, 7)[:-1]
    verts = np.column_stack((radius * np.cos(angles_v), radius * np.sin(angles_v)))

    # Sector fills
    sector_labels = [
        (r"S1: $D_A>D_B>D_C$", 30),
        (r"S2: $D_B>D_A>D_C$", 90),
        (r"S3: $D_B>D_C>D_A$", 150),
        (r"S4: $D_C>D_B>D_A$", 210),
        (r"S5: $D_C>D_A>D_B$", 270),
        (r"S6: $D_A>D_C>D_B$", 330),
    ]
    for i in range(6):
        v0 = verts[i]
        v1 = verts[(i + 1) % 6]
        tri = plt.Polygon(
            [np.array([0, 0]), v0, v1],
            closed=True,
            facecolor=SECTOR_COLORS[i],
            edgecolor="#aaa",
            lw=0.8,
        )
        ax.add_patch(tri)
        ang_rad = math.radians(sector_labels[i][1])
        lx = 0.62 * np.cos(ang_rad)
        ly = 0.62 * np.sin(ang_rad)
        ax.text(
            lx,
            ly,
            sector_labels[i][0],
            ha="center",
            va="center",
            fontsize=7.2,
            fontweight="bold",
        )

    # Hexagon outline
    poly = np.vstack((verts, verts[0]))
    ax.plot(poly[:, 0], poly[:, 1], "k-", lw=2)

    # Vertex labels (space vectors)
    sv_labels = [
        "$V_1(100)$",
        "$V_2(110)$",
        "$V_3(010)$",
        "$V_4(011)$",
        "$V_5(001)$",
        "$V_6(101)$",
    ]
    for i, (v, lbl) in enumerate(zip(verts, sv_labels)):
        ax.annotate(
            "",
            xy=v * 1.02,
            xytext=(0, 0),
            arrowprops=dict(arrowstyle="-|>", color="#444", lw=1.5),
        )
        offset = v * 1.18
        ax.text(
            offset[0],
            offset[1],
            lbl,
            ha="center",
            va="center",
            fontsize=8.5,
            color="#333",
        )

    # Reference vector
    ref_angle = math.radians(45)
    ref = np.array([0.65 * np.cos(ref_angle), 0.65 * np.sin(ref_angle)])
    ax.annotate(
        "",
        xy=ref,
        xytext=(0, 0),
        arrowprops=dict(arrowstyle="-|>", color="red", lw=2.5),
    )
    ax.text(
        ref[0] + 0.05,
        ref[1] + 0.05,
        r"$\vec{V}_{ref}$",
        color="red",
        fontsize=11,
        fontweight="bold",
    )

    ax.set_xlim(-1.45, 1.45)
    ax.set_ylim(-1.45, 1.45)
    ax.axhline(0, color="gray", lw=0.5, alpha=0.5)
    ax.axvline(0, color="gray", lw=0.5, alpha=0.5)
    ax.set_xlabel(r"$\alpha$ axis", fontsize=10)
    ax.set_ylabel(r"$\beta$ axis", fontsize=10)
    ax.set_title(
        "Fig. 4 – SVM Hexagon: Six Sectors with Duty Cycle Ordering\n"
        "and Active Space Vectors",
        fontsize=10,
        pad=10,
    )
    fig.tight_layout()
    return fig


def fig_blind_zone() -> plt.Figure:
    """Show window width variation vs electrical angle — blind zones near boundaries."""
    fig, axes = plt.subplots(2, 1, figsize=(8.5, 5.5), sharex=True)

    theta_deg = np.linspace(0, 360, 3600)
    theta = np.radians(theta_deg)

    # SVM reference duty cycles
    va = np.sin(theta)
    vb = np.sin(theta - 2 * np.pi / 3)
    vc = np.sin(theta + 2 * np.pi / 3)
    vmax = np.maximum(np.maximum(va, vb), vc)
    vmin = np.minimum(np.minimum(va, vb), vc)
    ucm = 0.5 * (vmax + vmin)
    va -= ucm
    vb -= ucm
    vc -= ucm

    # Normalised duty cycles [0, 1]
    Da = 0.5 * (va + 1)
    Db = 0.5 * (vb + 1)
    Dc = 0.5 * (vc + 1)

    MI = 0.866  # SVM modulation index (linear range)
    Da = 0.5 + 0.5 * MI * (va / np.max(np.abs(va)))
    Db = 0.5 + 0.5 * MI * (vb / np.max(np.abs(vb)))
    Dc = 0.5 + 0.5 * MI * (vc / np.max(np.abs(vc)))

    # Re-compute correctly
    Da = np.clip(0.5 * (va + 1), 0, 1)
    Db = np.clip(0.5 * (vb + 1), 0, 1)
    Dc = np.clip(0.5 * (vc + 1), 0, 1)

    T_half = 50.0  # µs
    dead = 2.5

    stacked = np.vstack([Da, Db, Dc])
    Dmax = np.max(stacked, axis=0)
    Dmid = np.median(stacked, axis=0)
    Dmin = np.min(stacked, axis=0)

    W1 = T_half * (Dmax - Dmid)
    W2 = T_half * (Dmid - Dmin)
    W1_eff = np.maximum(W1 - dead, 0)
    W2_eff = np.maximum(W2 - dead, 0)

    ax0, ax1 = axes

    ax0.fill_between(
        theta_deg,
        W1_eff,
        0,
        where=W1_eff > 1.5,
        color="#00aa44",
        alpha=0.5,
        label="Observable $W_1$",
    )
    ax0.fill_between(
        theta_deg,
        W1_eff,
        0,
        where=W1_eff <= 1.5,
        color="#e04040",
        alpha=0.5,
        label="Blind $W_1$",
    )
    ax0.plot(theta_deg, W1, "g--", lw=1, alpha=0.7, label=r"$W_1$ (ideal)")
    ax0.axhline(dead, color="k", lw=1, ls=":", alpha=0.8, label=r"$t_d = 2.5\,\mu s$")
    ax0.axhline(
        1.5, color="purple", lw=1, ls="-.", alpha=0.8, label=r"$t_{acq,min}=1.5\,\mu s$"
    )
    ax0.set_ylabel(r"$W_1$ ($\mu$s)", fontsize=10)
    ax0.legend(fontsize=8, ncol=2)
    ax0.grid(True, alpha=0.3)

    ax1.fill_between(
        theta_deg,
        W2_eff,
        0,
        where=W2_eff > 1.5,
        color="#aa5500",
        alpha=0.5,
        label="Observable $W_2$",
    )
    ax1.fill_between(
        theta_deg,
        W2_eff,
        0,
        where=W2_eff <= 1.5,
        color="#e04040",
        alpha=0.5,
        label="Blind $W_2$",
    )
    ax1.plot(theta_deg, W2, "orange", lw=1, ls="--", alpha=0.7, label=r"$W_2$ (ideal)")
    ax1.axhline(dead, color="k", lw=1, ls=":", alpha=0.8, label=r"$t_d = 2.5\,\mu s$")
    ax1.axhline(1.5, color="purple", lw=1, ls="-.", alpha=0.8, label=r"$t_{acq,min}$")
    ax1.set_ylabel(r"$W_2$ ($\mu$s)", fontsize=10)
    ax1.legend(fontsize=8, ncol=2)
    ax1.grid(True, alpha=0.3)

    # Sector boundaries
    for ax in axes:
        for ang in [0, 60, 120, 180, 240, 300, 360]:
            ax.axvline(ang, color="gray", lw=0.8, ls="--", alpha=0.5)
        for i, ang in enumerate([30, 90, 150, 210, 270, 330]):
            ax.text(
                ang,
                ax.get_ylim()[1] * 0.92 if ax.get_ylim()[1] > 0 else 1,
                f"S{i + 1}",
                ha="center",
                fontsize=8,
                color="gray",
            )

    ax1.set_xlabel("Electrical angle (degrees)", fontsize=10)
    ax1.set_xlim(0, 360)
    ax1.set_xticks([0, 60, 120, 180, 240, 300, 360])

    fig.suptitle(
        r"Fig. 5 – Effective Acquisition Window Widths $W_1^{eff}$ and $W_2^{eff}$ over One Electrical Cycle"
        "\n"
        r"(SVM, $T_{half}=50\,\mu s$, $t_d=2.5\,\mu s$). Red: blind zones near sector boundaries.",
        fontsize=9.5,
    )
    fig.tight_layout()
    return fig


def fig_edge_aligned_shift() -> plt.Figure:
    """Left-aligned PWM with and without phase shift — showing window creation."""
    fig, axes = plt.subplots(3, 2, figsize=(9, 6), sharey="row")
    fig.subplots_adjust(hspace=0.4, wspace=0.3)

    T = 1.0
    t = np.linspace(0, T, 2000)
    Da, Db, Dc = 0.75, 0.55, 0.30

    # Left-aligned (no shift)
    def left_pulse(D):
        return np.where(t < D * T, 1.0, 0.0)

    def shifted_left_pulse(D, shift):
        t_shifted = t - shift
        return np.where((t_shifted >= 0) & (t_shifted < D * T), 1.0, 0.0)

    # Optimal shift for left-aligned: stagger so windows open
    # Shift B by T*(Da-Db)/2, shift C by T*(Da-Dc)/2
    shift_b = T * (Da - Db) / 2
    shift_c = T * (Da - Dc) / 2
    # Simplified shift strategy: divide T evenly among phases
    shift_b_simple = T / 6
    shift_c_simple = T / 3

    scenarios = [
        (
            t,
            left_pulse(Da),
            left_pulse(Db),
            left_pulse(Dc),
            "No phase shift\n(left-aligned standard)",
        ),
        (
            t,
            shifted_left_pulse(Da, 0),
            shifted_left_pulse(Db, shift_b_simple),
            shifted_left_pulse(Dc, shift_c_simple),
            f"Phase shift applied\n(ΔtB={shift_b_simple:.2f}T, ΔtC={shift_c_simple:.2f}T)",
        ),
    ]

    for col, (t_arr, pa, pb, pc, title) in enumerate(scenarios):
        for row, (pulse, phase, col_hex) in enumerate(
            zip(
                [pa, pb, pc],
                ["A", "B", "C"],
                [PHASE_COLORS["A"], PHASE_COLORS["B"], PHASE_COLORS["C"]],
            )
        ):
            ax = axes[row, col]
            ax.fill_between(t_arr, 0, pulse, alpha=0.35, color=col_hex)
            ax.step(t_arr, pulse, color=col_hex, lw=1.5, where="post")
            ax.set_ylim(-0.2, 1.4)
            ax.set_yticks([0, 1])
            ax.grid(True, alpha=0.3)
            if row == 0:
                ax.set_title(title, fontsize=9)
            if col == 0:
                ax.set_ylabel(f"Phase {phase}", fontsize=9)

        # Annotate window for shifted case
        if col == 1:
            # Find where only A is on and B transitions
            pass

    for ax in axes[-1]:
        ax.set_xlabel("Time (t / T)", fontsize=9)

    fig.suptitle(
        "Fig. 6 – Left-Aligned (Edge-Aligned) PWM: Standard vs. Phase-Shifted Strategy\n"
        "Phase shifting creates isolated observation windows avoiding simultaneous transitions",
        fontsize=9.5,
    )
    return fig


def fig_dpwm_effect() -> plt.Figure:
    """DPWM120 clamping effect on acquisition windows."""
    fig, axes = plt.subplots(4, 1, figsize=(8.5, 6), sharex=True)
    fig.subplots_adjust(hspace=0.1, left=0.13, right=0.97, top=0.90, bottom=0.09)

    theta_deg = np.linspace(0, 360, 3600)
    theta = np.radians(theta_deg)

    va = np.sin(theta)
    vb = np.sin(theta - 2 * np.pi / 3)
    vc = np.sin(theta + 2 * np.pi / 3)
    vmax = np.maximum(np.maximum(va, vb), vc)
    ucm_svm = -0.5 * (
        np.maximum(np.maximum(va, vb), vc) + np.minimum(np.minimum(va, vb), vc)
    )

    # DPWM120 MAX: clamp max phase to +1
    tas = 0.5 * (va + 1)
    tbs = 0.5 * (vb + 1)
    tcs = 0.5 * (vc + 1)
    tmax = np.maximum(np.maximum(tas, tbs), tcs)
    toffset_dpwm_max = 1.0 - tmax
    Da_dpwm = np.clip(tas + toffset_dpwm_max, 0, 1)
    Db_dpwm = np.clip(tbs + toffset_dpwm_max, 0, 1)
    Dc_dpwm = np.clip(tcs + toffset_dpwm_max, 0, 1)

    T_half = 50.0
    dead = 2.0

    stacked = np.vstack([Da_dpwm, Db_dpwm, Dc_dpwm])
    Dmax = np.max(stacked, axis=0)
    Dmid = np.median(stacked, axis=0)
    Dmin = np.min(stacked, axis=0)
    W1 = np.maximum(T_half * (Dmax - Dmid) - dead, 0)
    W2 = np.maximum(T_half * (Dmid - Dmin) - dead, 0)

    for ax, data, lbl, col in zip(
        axes[:3],
        [Da_dpwm, Db_dpwm, Dc_dpwm],
        ["$D_A$", "$D_B$", "$D_C$"],
        [PHASE_COLORS["A"], PHASE_COLORS["B"], PHASE_COLORS["C"]],
    ):
        ax.plot(theta_deg, data, color=col, lw=1.5, label=lbl)
        ax.fill_between(theta_deg, data, 0, alpha=0.15, color=col)
        # Highlight clamp regions
        ax.fill_between(
            theta_deg,
            1.0,
            0.98,
            where=data > 0.995,
            color="#e04040",
            alpha=0.5,
            label="Clamped to 1",
        )
        ax.fill_between(
            theta_deg,
            0.02,
            0.0,
            where=data < 0.005,
            color="#4040e0",
            alpha=0.5,
            label="Clamped to 0",
        )
        ax.set_ylim(-0.05, 1.1)
        ax.set_yticks([0, 0.5, 1])
        ax.set_ylabel(lbl, fontsize=9, rotation=0, labelpad=30)
        ax.grid(True, alpha=0.3)

    axes[-1].fill_between(
        theta_deg, W1, 0, alpha=0.5, color="#00aa44", label=r"$W_1^{eff}$"
    )
    axes[-1].fill_between(
        theta_deg, -W2, 0, alpha=0.5, color="#aa5500", label=r"$W_2^{eff}$"
    )
    axes[-1].axhline(0, color="k", lw=0.5)
    axes[-1].set_ylabel(r"$W^{eff}$ ($\mu$s)", fontsize=9, rotation=0, labelpad=42)
    axes[-1].legend(fontsize=8)
    axes[-1].grid(True, alpha=0.3)
    axes[-1].set_xlabel("Electrical angle (degrees)", fontsize=10)

    for ax in axes:
        for ang in [0, 60, 120, 180, 240, 300, 360]:
            ax.axvline(ang, color="gray", lw=0.7, ls="--", alpha=0.4)

    axes[-1].set_xlim(0, 360)
    axes[-1].set_xticks([0, 60, 120, 180, 240, 300, 360])

    fig.suptitle(
        "Fig. 7 – DPWM120 (MAX) Duty Cycle Profiles and Resulting Effective Acquisition Windows\n"
        "Red: clamped-to-1 region (no switching); Blue: clamped-to-0 region",
        fontsize=9.5,
    )
    return fig


def fig_compensation_strategies() -> plt.Figure:
    """Comparison of compensation strategies near a blind zone."""
    fig, axes = plt.subplots(1, 3, figsize=(9.5, 3.8))
    fig.subplots_adjust(wspace=0.35, top=0.82, bottom=0.15)

    T_half = 50.0
    dead = 2.0
    adc_min = 1.5
    min_w = dead + adc_min

    # Near-sector boundary: Da≈Db (W1 → 0)
    Da, Db, Dc = 0.62, 0.60, 0.30

    strategies = [
        ("Original\n(W1 blind)", Da, Db, Dc),
        ("Min-pulse\ninsertion", Da, max(Db, Da - (min_w / T_half)), Dc),
        (
            "Duty-cycle\nredistribution",
            Da + min_w / (2 * T_half),
            Db - min_w / (2 * T_half),
            Dc,
        ),
    ]

    for ax, (title, da, db, dc, *_) in zip(axes, strategies):
        stacked = sorted([da, db, dc], reverse=True)
        dmax, dmid, dmin = stacked
        w1 = max(T_half * (dmax - dmid) - dead, 0)
        w2 = max(T_half * (dmid - dmin) - dead, 0)

        t = np.linspace(0, T_half, 1000)
        t_on_max = T_half * (1 - dmax)
        t_on_mid = T_half * (1 - dmid)
        t_on_min = T_half * (1 - dmin)

        obs_color_w1 = "#00aa44" if w1 >= adc_min else "#e04040"
        obs_color_w2 = "#aa5500" if w2 >= adc_min else "#e04040"

        ax.axvspan(
            t_on_min, t_on_min + dead, ymin=0.1, ymax=0.7, color="#e04040", alpha=0.6
        )
        ax.axvspan(
            t_on_min + dead, t_on_mid, ymin=0.1, ymax=0.7, color=obs_color_w2, alpha=0.4
        )
        ax.axvspan(
            t_on_mid, t_on_mid + dead, ymin=0.1, ymax=0.7, color="#e04040", alpha=0.6
        )
        ax.axvspan(
            t_on_mid + dead, t_on_max, ymin=0.1, ymax=0.7, color=obs_color_w1, alpha=0.4
        )
        ax.axvspan(
            t_on_max, t_on_max + dead, ymin=0.1, ymax=0.7, color="#e04040", alpha=0.6
        )
        ax.axvspan(
            t_on_max + dead, T_half, ymin=0.1, ymax=0.7, color="#4040e0", alpha=0.15
        )

        ax.set_xlim(0, T_half)
        ax.set_ylim(0, 1)
        ax.set_yticks([])
        obs_str = f"W1={w1:.1f}µs {'✓' if w1 >= adc_min else '✗'}  W2={w2:.1f}µs {'✓' if w2 >= adc_min else '✗'}"
        ax.set_title(f"{title}\n{obs_str}", fontsize=8.5, pad=4)
        ax.set_xlabel(r"$t$ ($\mu$s)", fontsize=9)
        ax.grid(True, alpha=0.3)
        ax.axhline(0.4, color="k", lw=1, ls="-", alpha=0.5)

    # Legend
    patch_dead = mpatches.Patch(color="#e04040", alpha=0.6, label="Dead time")
    patch_obs = mpatches.Patch(color="#00aa44", alpha=0.5, label="Observable window")
    patch_blind = mpatches.Patch(color="#e04040", alpha=0.4, label="Blind window")
    patch_all = mpatches.Patch(color="#4040e0", alpha=0.2, label="All-ON region")
    fig.legend(
        handles=[patch_dead, patch_obs, patch_blind, patch_all],
        loc="lower center",
        ncol=4,
        fontsize=8.5,
        bbox_to_anchor=(0.5, 0.0),
    )

    fig.suptitle(
        "Fig. 8 – Blind Zone Compensation Strategies (Half-Period View)\n"
        "Near sector boundary: $D_A \\approx D_B = 0.62 / 0.60$",
        fontsize=9.5,
        y=0.95,
    )
    return fig


def fig_reconstruction_algorithm() -> plt.Figure:
    """Flowchart of the current reconstruction algorithm."""
    fig, ax = plt.subplots(figsize=(7, 8))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 14)
    ax.axis("off")
    ax.set_facecolor("white")

    def box(x, y, w, h, text, color="#d6e4f7", text_size=9):
        ax.add_patch(
            Rectangle(
                (x - w / 2, y - h / 2),
                w,
                h,
                facecolor=color,
                edgecolor="#444",
                lw=1.2,
                zorder=3,
            )
        )
        ax.text(
            x,
            y,
            text,
            ha="center",
            va="center",
            fontsize=text_size,
            wrap=True,
            zorder=4,
        )

    def diamond(x, y, w, h, text, color="#fff3cd"):
        dx, dy = w / 2, h / 2
        pts = [(x, y + dy), (x + dx, y), (x, y - dy), (x - dx, y)]
        ax.add_patch(
            plt.Polygon(
                pts, closed=True, facecolor=color, edgecolor="#444", lw=1.2, zorder=3
            )
        )
        ax.text(x, y, text, ha="center", va="center", fontsize=8.5, zorder=4)

    def arrow(x1, y1, x2, y2, label=""):
        ax.annotate(
            "",
            xy=(x2, y2),
            xytext=(x1, y1),
            arrowprops=dict(arrowstyle="-|>", color="#333", lw=1.5),
            zorder=5,
        )
        if label:
            ax.text(
                (x1 + x2) / 2 + 0.15, (y1 + y2) / 2, label, fontsize=8, color="#555"
            )

    # Blocks
    box(5, 13, 6, 0.9, "START: New PWM period begins", color="#d0f0d0")
    arrow(5, 12.55, 5, 11.9)
    box(5, 11.4, 7, 0.9, "Read duty cycles: $D_A$, $D_B$, $D_C$\nfrom modulator output")
    arrow(5, 10.95, 5, 10.25)
    box(
        5,
        9.75,
        7,
        0.9,
        "Sort: $(D_{max}, D_{mid}, D_{min})$\n→ identify phase ordering",
    )
    arrow(5, 9.30, 5, 8.6)
    box(
        5,
        8.1,
        7,
        0.9,
        r"Compute: $W_1=\frac{T}{2}(D_{max}-D_{mid})$,"
        "\n" + r"$W_2=\frac{T}{2}(D_{mid}-D_{min})$",
    )
    arrow(5, 7.65, 5, 7.0)
    diamond(5, 6.5, 6, 0.9, r"$W_1^{eff} = W_1 - t_d > t_{acq,min}$?")
    arrow(5, 6.05, 5, 5.4)
    ax.text(5.1, 5.7, "Yes", fontsize=8, color="green")
    arrow(8.0, 6.5, 9.5, 6.5)
    ax.text(8.8, 6.6, "No", fontsize=8, color="red")
    box(
        5,
        4.9,
        6,
        0.9,
        r"Sample shunt at center of $W_1$"
        "\n" + r"→ $\hat{I}_{max} = v_{shunt} / R_{shunt}$",
        color="#e8f4fd",
    )
    arrow(5, 4.45, 5, 3.8)
    diamond(5, 3.3, 6, 0.9, r"$W_2^{eff} = W_2 - t_d > t_{acq,min}$?")
    arrow(5, 2.85, 5, 2.2)
    ax.text(5.1, 2.5, "Yes", fontsize=8, color="green")
    box(
        5,
        1.7,
        6,
        0.9,
        r"$\hat{I}_{min} = -v_{shunt}(W_2) / R_{shunt}$"
        "\n" + r"$\hat{I}_{mid} = -\hat{I}_{max} - \hat{I}_{min}$ (KCL)",
        color="#e8f4fd",
    )
    arrow(5, 1.25, 5, 0.6)
    box(
        5,
        0.3,
        6,
        0.5,
        r"Output: $\hat{I}_A$, $\hat{I}_B$, $\hat{I}_C$",
        color="#d0f0d0",
    )

    # Compensation path
    box(9.5, 5.5, 1.5, 0.6, "Apply\ncompensation", color="#ffe0e0", text_size=8)

    ax.set_title(
        "Fig. 9 – Single Shunt Current Reconstruction Algorithm Flowchart",
        fontsize=10,
        pad=10,
    )
    fig.tight_layout()
    return fig


def fig_symbol_table_figure() -> plt.Figure:
    """Symbol reference table as a figure for embedding."""
    fig, ax = plt.subplots(figsize=(8.5, 4.5))
    ax.axis("off")
    symbols = [
        (r"$T_{PWM}$", "PWM carrier period", "s"),
        (r"$f_{PWM}$", "PWM carrier frequency", "Hz"),
        (r"$D_x$ (x = A,B,C)", "Per-phase duty cycle (0 to 1)", "—"),
        (r"$D_{max}, D_{mid}, D_{min}$", "Sorted per-period duty cycles", "—"),
        (r"$t_{on,x}$", "Turn-on time of phase x (center-aligned)", "s"),
        (r"$W_1$", "Ideal acquisition window 1 width", r"$\mu$s"),
        (r"$W_2$", "Ideal acquisition window 2 width", r"$\mu$s"),
        (
            r"$W_1^{eff}, W_2^{eff}$",
            "Effective window widths (after dead-time reduction)",
            r"$\mu$s",
        ),
        (r"$t_d$", "Dead time (blanking time)", r"$\mu$s"),
        (r"$t_{acq,min}$", "Minimum ADC acquisition time", r"$\mu$s"),
        (r"$R_{shunt}$", "DC-bus shunt resistor value", r"$\Omega$"),
        (r"$i_{dc}$", "Instantaneous DC-bus shunt current", "A"),
        (r"$\hat{I}_x$", "Reconstructed phase-current estimate for phase x", "A"),
        (r"$V_{DC}$", "DC-link voltage", "V"),
        (r"$MI$", "Modulation index", "—"),
        (r"$\theta_e$", "Electrical angle of the reference vector", "rad"),
    ]
    col_labels = ["Symbol", "Description", "Unit"]
    rows = [[s[0], s[1], s[2]] for s in symbols]
    tbl = ax.table(
        cellText=rows,
        colLabels=col_labels,
        loc="center",
        cellLoc="left",
    )
    tbl.auto_set_font_size(False)
    tbl.set_fontsize(9)
    tbl.auto_set_column_width(col=[0, 1, 2])
    ax.set_title("Table I – Symbol Glossary", fontsize=11, pad=15)
    fig.tight_layout()
    return fig


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------


def build_document(out_path: str) -> None:
    doc = Document()

    # ---------- Page setup ----------
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # ---------- Title ----------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(
        "Single Shunt Current Reconstruction in\n"
        "Three-Phase PWM Inverters:\n"
        "Theory, Modulation Strategies, Acquisition Window Analysis,\n"
        "and Blind-Zone Compensation Methods"
    )
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()
    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_p.add_run("Technical Report — SVM Analyst Project\n").italic = True
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("April 2026").italic = True

    doc.add_paragraph()

    # ---------- Abstract ----------
    _heading(doc, "Abstract", 1)
    doc.add_paragraph(
        "This report presents a comprehensive analysis of single shunt current reconstruction "
        "(SSCR) techniques for three-phase voltage source inverters (VSI). The single shunt "
        "topology is attractive for cost-sensitive drives because one resistor placed in the "
        "DC-bus negative rail can provide all three phase currents at the expense of specific "
        "PWM timing constraints. This document establishes the mathematical foundation of the "
        "acquisition window algorithm, examines the dependency of measurable window widths on "
        "the duty cycles produced by different modulation strategies—including Sinusoidal PWM "
        "(SPWM), Third-Harmonic Injection PWM (THIPWM), Space Vector Modulation (SVM), and "
        "all six Discontinuous PWM (DPWM) variants—and derives the conditions under which "
        "reconstruction fails (blind zones). Compensation methods for center-aligned and "
        "edge-aligned PWM schemes are presented, including minimum-pulse insertion, "
        "duty-cycle redistribution, and phase-time shifting. The analysis is directly tied "
        "to the SVM Analyst simulation platform."
    )

    # ---------- Keywords ----------
    kw_p = doc.add_paragraph()
    _bold(kw_p, "Keywords: ")
    kw_p.add_run(
        "single shunt current sensing, PWM inverter, space vector modulation, DPWM, "
        "blind zone, dead time, ADC acquisition window, current reconstruction, "
        "motor control, IGBT inverter."
    )

    doc.add_page_break()

    # ---------- Section I: Introduction ----------
    _heading(doc, "I. Introduction", 1)
    doc.add_paragraph(
        "Accurate measurement of three-phase motor currents is a fundamental requirement "
        "for field-oriented control (FOC) and direct torque control (DTC) of permanent-magnet "
        "synchronous motors (PMSM) and induction motors (IM) [1]. Traditionally, dedicated "
        "current sensors (Hall-effect or shunt-based) are placed in each motor phase, "
        "requiring three hardware channels. This approach is accurate and robust but increases "
        "bill-of-materials cost and PCB complexity."
    )
    doc.add_paragraph(
        "The single shunt current reconstruction technique reduces the sensor count to one "
        "resistor in the DC-bus negative rail [2][3]. The shunt carries the sum of all "
        "active-low-side device currents and, during specific time intervals within every "
        "PWM period, carries exactly one phase current at a time. By careful ADC triggering "
        "within those intervals, all three phase currents can be estimated from two successive "
        "samples per period."
    )
    doc.add_paragraph(
        "This approach is well established in low-cost drive inverters but presents specific "
        "challenges: (i) near sector boundaries of the modulation pattern the acquisition "
        "interval narrows and may vanish (blind zones); (ii) parasitic ringing after IGBT "
        "switching mandates a minimum settling time before sampling; (iii) dead-time blanking "
        "further reduces the usable window; and (iv) different modulation strategies "
        "generate markedly different window-width profiles across the electrical cycle [4][5]."
    )
    doc.add_paragraph(
        "Section II reviews the inverter topology and DC-bus shunt principle. Section III "
        "develops the PWM timing equations for center-aligned carriers. Section IV extends "
        "the analysis to all SVM sectors and all modulation flavors available in the SVM "
        "Analyst project. Section V characterizes blind zones. Section VI presents "
        "compensation strategies. Section VII addresses edge-aligned PWM and phase-shift "
        "methods. Section VIII discusses implementation in the SVM Analyst tool. Conclusions "
        "follow in Section IX."
    )

    # ---------- Section II: Inverter Topology ----------
    _heading(doc, "II. Three-Phase VSI with Single DC-Bus Shunt", 1)
    doc.add_paragraph(
        "A three-phase two-level voltage source inverter (VSI) consists of three half-bridge "
        "legs, each containing a high-side switch "
        r"(Q_xH, x ∈ {A, B, C}) and a low-side switch (Q_xL) "
        "as shown in Fig. 1. Dead time "
        r"(t_d) "
        "is inserted around each switching transition to prevent cross-conduction."
    )
    doc.add_paragraph(
        "When a low-side switch Q_xL conducts, the corresponding phase current i_x flows "
        "through the DC-bus negative rail and through the shunt resistor R_shunt. The "
        "instantaneous DC-bus current is therefore:"
    )
    eq_p = doc.add_paragraph()
    eq_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq_p.add_run("i_dc(t) = Σ [S_xL(t) · i_x(t)]  for x ∈ {A, B, C}  … (1)")
    doc.add_paragraph(
        "where S_xL ∈ {0, 1} is the logical state of the low-side switch for phase x. "
        "The shunt voltage is:"
    )
    eq_p2 = doc.add_paragraph()
    eq_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq_p2.add_run("v_shunt(t) = R_shunt · i_dc(t)  … (2)")

    buf = _fig_to_stream(fig_inverter_topology())
    doc.add_picture(buf, width=Inches(5.8))
    _add_caption(
        doc, "Fig. 1 – Three-Phase Half-Bridge VSI with Single DC-Bus Shunt Resistor."
    )

    # ---------- Section III: Center-Aligned PWM Timing ----------
    _heading(doc, "III. Center-Aligned PWM: Acquisition Window Formation", 1)

    _heading(doc, "III-A. Carrier Comparison and Duty Cycle", 2)
    doc.add_paragraph(
        "In center-aligned (symmetric, up-down counter) PWM, each phase x has a duty cycle "
        "D_x ∈ [0, 1] controlling the fraction of the period T_PWM during which Q_xH is ON "
        "and Q_xL is OFF. The turn-on time of phase x within a half-period is:"
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("t_on,x = (T_PWM / 2) · (1 − D_x)  … (3)")
    doc.add_paragraph(
        "Because Q_xL is the complement of Q_xH (with dead time enforced), Q_xL is ON "
        "precisely when Q_xH is OFF, i.e., during [0, t_on,x] and [T_PWM − t_on,x, T_PWM]. "
        "Therefore, i_x appears on the shunt during those intervals."
    )

    _heading(doc, "III-B. Duty Cycle Ordering and Window Identification", 2)
    doc.add_paragraph("For any PWM period, the three duty cycles are sorted:")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("D_max ≥ D_mid ≥ D_min  (phases reassigned accordingly)  … (4)")
    doc.add_paragraph(
        "Using (3), the corresponding turn-on times satisfy t_on,max ≤ t_on,mid ≤ t_on,min. "
        "In the first half of the PWM period, the following state sequence occurs:"
    )

    # State table
    tbl = doc.add_table(rows=5, cols=3)
    tbl.style = "Table Grid"
    hdrs = ["Time interval", "Low-side switches ON", "i_dc (shunt current)"]
    for j, h in enumerate(hdrs):
        tbl.rows[0].cells[j].text = h
        tbl.rows[0].cells[j].paragraphs[0].runs[0].bold = True
        _shade_cell(tbl.rows[0].cells[j], "D6E4F7")
    data = [
        ("[0,  t_on,min]", "All three", "≈ 0  (balanced, zero vector)"),
        ("[t_on,min,  t_on,mid]  ← W₂", "max + mid only", "= −i_min  ← Sample 2"),
        ("[t_on,mid,  t_on,max]  ← W₁", "max only", "= i_max  ← Sample 1"),
        ("[t_on,max,  T/2]", "None (all high-side ON)", "≈ 0  (zero vector)"),
    ]
    for i, row_data in enumerate(data):
        for j, val in enumerate(row_data):
            tbl.rows[i + 1].cells[j].text = val

    doc.add_paragraph()
    _add_caption(
        doc, "Table II – DC-Bus Shunt Current States in Center-Aligned PWM Half-Period."
    )

    doc.add_paragraph(
        "Window W₁ corresponds to the interval when only the phase with the highest duty "
        "cycle has its low-side device ON, providing a direct measurement of i_max. "
        "Window W₂ gives i_dc = i_max + i_mid = −i_min (by Kirchhoff's current law, "
        "i_A + i_B + i_C = 0 for star-connected balanced load)."
    )

    _heading(doc, "III-C. Window Width Equations", 2)
    doc.add_paragraph("The ideal window widths within one half-period are:")
    for eq, lbl in [
        ("W₁ = (T_PWM / 2) · (D_max − D_mid)  … (5)", ""),
        ("W₂ = (T_PWM / 2) · (D_mid − D_min)  … (6)", ""),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(eq)

    doc.add_paragraph(
        "After subtracting the dead time t_d (during which the shunt current is undefined "
        "due to switch-state transients), the effective widths are:"
    )
    for eq in ["W₁_eff = max(W₁ − t_d, 0)  … (7)", "W₂_eff = max(W₂ − t_d, 0)  … (8)"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(eq)

    doc.add_paragraph(
        "Reconstruction is possible for a given window when the effective width exceeds a "
        "minimum ADC acquisition time t_acq,min:"
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("W_eff > t_acq,min  (reconstruction condition)  … (9)")

    buf2 = _fig_to_stream(fig_center_aligned_pwm())
    doc.add_picture(buf2, width=Inches(6.0))
    _add_caption(
        doc,
        "Fig. 2 – Center-Aligned PWM Timing Diagram (Sector 1, D_A > D_B > D_C). "
        "Green shading: W₁ (sample I_A); brown shading: W₂ (sample −I_C).",
    )

    buf3 = _fig_to_stream(fig_window_geometry())
    doc.add_picture(buf3, width=Inches(5.5))
    _add_caption(
        doc,
        "Fig. 3 – Half-Period Zoom: Dead-Time Consumption and Effective Acquisition "
        "Windows (T_half = 50 µs, t_d = 2 µs).",
    )

    _heading(doc, "III-D. Current Reconstruction Identities", 2)
    doc.add_paragraph(
        "From the two samples, all three phase currents are obtained as follows. Let Sample 1 "
        "be taken at the center of W₁ and Sample 2 at the center of W₂:"
    )
    for eq in [
        "î_max = Sample₁  … (10)",
        "î_min = −Sample₂  … (11)",
        "î_mid = −î_max − î_min  (KCL)  … (12)",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(eq)
    doc.add_paragraph(
        "The identities (10)–(12) are valid for all six SVM sectors because the "
        "window labeling always follows the sorted duty-cycle ordering of (4), "
        "regardless of which physical phase (A, B, or C) is D_max, D_mid, or D_min."
    )

    # ---------- Section IV: SVM Sector Analysis ----------
    _heading(doc, "IV. SVM Sector Analysis and Phase Assignment", 1)

    _heading(doc, "IV-A. SVM Hexagon and Sector Definition", 2)
    doc.add_paragraph(
        "Space Vector Modulation represents the three-phase voltage reference as a rotating "
        "voltage vector V_ref in the stationary αβ frame. The complex plane is partitioned "
        "into six 60° sectors, delimited by the six active voltage vectors of the inverter "
        "(Fig. 4). In each sector, the duty cycles produced by the SVM algorithm follow a "
        "known ordering of the three phases."
    )

    buf4 = _fig_to_stream(fig_svm_sectors())
    doc.add_picture(buf4, width=Inches(4.0))
    _add_caption(
        doc,
        "Fig. 4 – SVM Hexagon with Six Sectors and Their Corresponding "
        "Duty-Cycle Phase Orderings.",
    )

    _heading(doc, "IV-B. Per-Sector Phase Ordering Table", 2)
    doc.add_paragraph(
        "Table III summarises the SVM sector boundaries, the duty-cycle ordering, and "
        "which physical phase current is measured by W₁ and W₂ in each sector."
    )
    tbl2 = doc.add_table(rows=8, cols=5)
    tbl2.style = "Table Grid"
    tbl2_hdrs = ["Sector", "θ_e range", "D ordering", "W₁ → measures", "W₂ → measures"]
    for j, h in enumerate(tbl2_hdrs):
        tbl2.rows[0].cells[j].text = h
        tbl2.rows[0].cells[j].paragraphs[0].runs[0].bold = True
        _shade_cell(tbl2.rows[0].cells[j], "D6E4F7")
    sector_data = [
        ("1", "0° – 60°", "D_A > D_B > D_C", "+I_A", "−I_C"),
        ("2", "60° – 120°", "D_B > D_A > D_C", "+I_B", "−I_C"),
        ("3", "120° – 180°", "D_B > D_C > D_A", "+I_B", "−I_A"),
        ("4", "180° – 240°", "D_C > D_B > D_A", "+I_C", "−I_A"),
        ("5", "240° – 300°", "D_C > D_A > D_B", "+I_C", "−I_B"),
        ("6", "300° – 360°", "D_A > D_C > D_B", "+I_A", "−I_B"),
    ]
    for i, row_vals in enumerate(sector_data):
        for j, val in enumerate(row_vals):
            tbl2.rows[i + 1].cells[j].text = val
    doc.add_paragraph()
    _add_caption(
        doc,
        "Table III – SVM Sector Duty-Cycle Ordering and Corresponding "
        "Shunt Measurements (W₁ and W₂).",
    )

    _heading(doc, "IV-C. Extension to All Modulation Modes", 2)
    doc.add_paragraph(
        "The acquisition window algorithm of (5)–(9) is modulation-agnostic: it requires "
        "only the instantaneous per-period duty cycles D_A, D_B, D_C. However, different "
        "modulation strategies produce different duty-cycle profiles and therefore different "
        "window-width distributions over the electrical cycle."
    )

    tbl3 = doc.add_table(rows=9, cols=3)
    tbl3.style = "Table Grid"
    tbl3_hdrs = [
        "Modulation Mode",
        "Common-Mode Injected",
        "Window-Width Characteristic",
    ]
    for j, h in enumerate(tbl3_hdrs):
        tbl3.rows[0].cells[j].text = h
        tbl3.rows[0].cells[j].paragraphs[0].runs[0].bold = True
        _shade_cell(tbl3.rows[0].cells[j], "D6E4F7")
    mod_data = [
        ("SPWM (Sinusoidal)", "None", "Smaller windows; D differs by up to 0.5·MI"),
        ("THIPWM 1/6", "1/6 × sin(3θ)", "Moderate windows; 15% bus utilisation gain"),
        (
            "THIPWM 1/4",
            "1/4 × sin(3θ)",
            "Near-SVM performance at explicit injection cost",
        ),
        (
            "Custom THIPWM",
            "k × sin(3θ), k adjustable",
            "Windows scale with injection factor k",
        ),
        ("SVM", "Vmin/2 + Vmax/2", "Widest windows; optimal for SSCR (reference)"),
        (
            "DPWM120 MAX",
            "Clamp max to 1 for 120°",
            "W₁ or W₂ collapses near clamp edge",
        ),
        (
            "DPWM120 MIN",
            "Clamp min to 0 for 120°",
            "Similar to MAX but at D=0 boundary",
        ),
        (
            "DPWM60 / DPWM30",
            "Alternating clamp strategies",
            "Narrower but distributed blind zones",
        ),
    ]
    for i, row_vals in enumerate(mod_data):
        for j, val in enumerate(row_vals):
            tbl3.rows[i + 1].cells[j].text = val
    doc.add_paragraph()
    _add_caption(
        doc, "Table IV – Modulation Modes and Their Impact on SSCR Acquisition Windows."
    )

    buf5 = _fig_to_stream(fig_dpwm_effect())
    doc.add_picture(buf5, width=Inches(5.8))
    _add_caption(
        doc,
        "Fig. 7 – DPWM120 (MAX) Duty Cycle Profiles and Effective Acquisition "
        "Windows. Red: clamped-to-1 region; Blue: clamped-to-0 region.",
    )

    # ---------- Section V: Blind Zones ----------
    _heading(doc, "V. Blind Zones: Analysis and Characterisation", 1)
    doc.add_paragraph(
        "Blind zones arise whenever one or both effective window widths fall below t_acq,min. "
        "Two distinct mechanisms produce blind zones:"
    )
    doc.add_paragraph(
        "1. Sector boundary transition: As θ_e approaches a 60° boundary, the reference "
        "vector rotates toward the adjacent active vector. The two phases with closest duty "
        "cycles exchange their ordering (D_max ↔ D_mid or D_mid ↔ D_min), causing W₁ or W₂ "
        "to shrink to zero and then grow again in the new sector. The angular width of the "
        "blind zone depends on the modulation index MI and the PWM frequency ratio N = f_PWM / f_e."
    )
    doc.add_paragraph(
        "2. Clamped-phase effect (DPWM modes): When one phase is clamped to D = 1 (or D = 0), "
        "its turn-on time is at the very beginning (or end) of the half-period. This forces "
        "W₁ or W₂ to collapse to zero for the entire clamped region (typically 60° or 120°)."
    )
    doc.add_paragraph(
        "The angular extent of a sector-boundary blind zone Δθ_blind for W₁ is (for SVM):"
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(
        "Δθ_blind ≈ (2 / (π · MI)) · arcsin( (t_d + t_acq,min) / (T_PWM / 2) )  … (13)"
    )
    doc.add_paragraph(
        "For a typical drive: T_PWM = 100 µs (10 kHz), t_d = 2.5 µs, t_acq,min = 1.5 µs, "
        "MI = 0.8, Δθ_blind ≈ 7.5° on each side of a sector boundary, for a total of ~15° "
        "per sector boundary (12 boundaries per full cycle → ~180° total out of 360°)."
    )

    buf6 = _fig_to_stream(fig_blind_zone())
    doc.add_picture(buf6, width=Inches(6.0))
    _add_caption(
        doc,
        "Fig. 5 – Effective Acquisition Window Widths W₁_eff (top) and W₂_eff "
        "(bottom) over One Full Electrical Cycle (SVM, T_half = 50 µs, t_d = 2.5 µs). "
        "Red areas indicate blind zones near sector boundaries.",
    )

    # ---------- Section VI: Compensation ----------
    _heading(doc, "VI. Blind Zone Compensation for Center-Aligned PWM", 1)

    _heading(doc, "VI-A. Strategy Overview", 2)
    doc.add_paragraph(
        "Three main compensation strategies exist for center-aligned PWM blind zones [6][7]:"
    )

    strategies_list = [
        (
            "Minimum pulse insertion",
            "Forces a minimum separation between consecutive switching events. When |D_max − D_mid| < W_min / T_half, "
            "the duty cycles are artificially spread: D_max is increased and/or D_mid is decreased by δ = W_min / T_half "
            "until the window is measurable. This introduces a small voltage error "
            "that must be compensated by a feed-forward dead-time correction term.",
        ),
        (
            "Duty-cycle redistribution",
            "Redistributes the separation deficit symmetrically: one half is added to D_max and the other "
            "half subtracted from D_mid, minimising the fundamental voltage error compared to unilateral insertion.",
        ),
        (
            "Phase-current estimation (hold-last-value)",
            "When neither W₁ nor W₂ is observable (deep blind zone), the last valid measurement is held "
            "and a model-based predictor (first-order Kalman filter or simple integrator) extrapolates the current. "
            "This is acceptable for a few consecutive PWM periods but degrades with longer blind zones.",
        ),
    ]
    for name, desc in strategies_list:
        p = doc.add_paragraph(style="List Number")
        _bold(p, f"{name}: ")
        p.add_run(desc)

    buf7 = _fig_to_stream(fig_compensation_strategies())
    doc.add_picture(buf7, width=Inches(5.8))
    _add_caption(
        doc,
        "Fig. 8 – Blind Zone Compensation Near Sector Boundary (D_A ≈ D_B). "
        "Left: original (W₁ blind). Centre: min-pulse insertion. "
        "Right: duty-cycle redistribution.",
    )

    _heading(doc, "VI-B. Voltage Error Analysis", 2)
    doc.add_paragraph(
        "Minimum-pulse insertion introduces a duty-cycle perturbation δ per phase. The resulting "
        "fundamental output voltage error is proportional to δ · V_DC / T_PWM. For δ = W_min / T_half "
        "with W_min = t_d + t_acq,min = 4 µs and T_half = 50 µs:"
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(
        "ΔV_output = δ · V_DC = (4 µs / 50 µs) · V_DC = 0.08 · V_DC  (maximally, at boundary)  … (14)"
    )
    doc.add_paragraph(
        "This error is localised to the narrow angular region near the sector boundary. "
        "Feed-forward correction inverts the perturbation on the next period, limiting steady-state "
        "current ripple to within the dead-time compensation error band."
    )

    # ---------- Section VII: Edge-Aligned PWM ----------
    _heading(doc, "VII. Edge-Aligned PWM and Phase-Shift Strategy", 1)

    _heading(doc, "VII-A. Left-Aligned PWM Timing", 2)
    doc.add_paragraph(
        "In left-aligned (edge-aligned up-count) PWM, all three phase pulses start at t = 0. "
        "The turn-off times are:"
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("t_off,x = D_x · T_PWM  (left-aligned)  … (15)")
    doc.add_paragraph(
        "All three switching events are clustered near the beginning of the period. "
        "For the shunt to capture two independent measurements, transitions must be "
        "sufficiently separated. In the sorted ordering, the timing structure is:"
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(
        "W₁_LA = T_PWM · (D_max − D_mid) − t_d  … (16)\n"
        "W₂_LA = T_PWM · (D_mid − D_min) − t_d  … (17)"
    )
    doc.add_paragraph(
        "Note that factors are T_PWM (not T_PWM / 2 as in center-aligned), because the "
        "asymmetric carrier has no symmetry benefit. This means edge-aligned PWM provides "
        "twice the nominal window width compared to center-aligned for the same duty "
        "cycle difference — but the windows occur only once per period, not twice."
    )

    _heading(doc, "VII-B. Right-Aligned PWM Timing", 2)
    doc.add_paragraph(
        "In right-aligned PWM, all pulses end at T_PWM. Turn-on times are:"
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("t_on,x = T_PWM · (1 − D_x)  (right-aligned)  … (18)")
    doc.add_paragraph(
        "The window equations remain identical in form to (16)–(17) because the relative "
        "separations between turn-on events are the same as for left-aligned turn-off events."
    )

    _heading(doc, "VII-C. Phase-Shift Strategy for Edge-Aligned PWM", 2)
    doc.add_paragraph(
        "A fundamental limitation of standard edge-aligned PWM is that all three switching "
        "transitions occur within a short burst at the start (or end) of the period, increasing "
        "EMI and potentially producing zero-window conditions when duty cycles are close. "
        "The phase-shift strategy addresses this by adding a static time offset Δt_x to each "
        "phase's PWM reference:"
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("t_off,x_shifted = D_x · T_PWM + Δt_x   (shifted PWM)  … (19)")
    doc.add_paragraph(
        "The shifts are chosen to guarantee a minimum separation between any two transitions. "
        "A common choice for three phases is:"
    )
    for eq in [
        "Δt_A = 0  … (constant reference)",
        "Δt_B = T_PWM / 6  … (120° interleave offset)",
        "Δt_C = T_PWM / 3  … (240° interleave offset)",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(eq + "  … (20a–c)")
        break
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Δt_A = 0,   Δt_B = T_PWM/6,   Δt_C = T_PWM/3   … (20)")
    doc.add_paragraph(
        "With this choice, the transitions of B and C are guaranteed to be separated from A "
        "by at least T_PWM / 6, creating observable windows even when D_A ≈ D_B ≈ D_C "
        "(the worst case for window-width collapse in standard edge-aligned PWM). "
        "The phase shifts introduce a fundamental output voltage phase offset that can be "
        "compensated by adjusting the SVPWM reference angle accordingly."
    )

    buf8 = _fig_to_stream(fig_edge_aligned_shift())
    doc.add_picture(buf8, width=Inches(6.0))
    _add_caption(
        doc,
        "Fig. 6 – Left-Aligned PWM: Standard (left) vs. Phase-Shifted Strategy "
        "(right). Inter-phase time offsets create isolated measurement windows.",
    )

    # ---------- Section VIII: SVM Analyst Implementation ----------
    _heading(doc, "VIII. Integration in the SVM Analyst Tool", 1)

    _heading(doc, "VIII-A. New Module Architecture", 2)
    doc.add_paragraph(
        "The SSCR feature is implemented as a new Python module svm_shaper/single_shunt.py "
        "containing pure numpy computation, plus a SingleShuntDialog class in gui.py. "
        "The module exposes the following public API:"
    )
    for name, desc in [
        (
            "compute_single_shunt_analysis(config, result)",
            "Computes the full per-period acquisition analysis from an existing SimulationResult.",
        ),
        (
            "get_sector_from_angle(theta_deg)",
            "Returns SVM sector (1–6) for a given electrical angle.",
        ),
        (
            "get_duty_ordering(d_a, d_b, d_c)",
            "Returns sorted (max, mid, min) with associated phase labels.",
        ),
        (
            "build_pwm_period_pulse_shapes(period, dead_time_us)",
            "Builds time-domain pulse shape arrays for a single PWM period, suitable for the zoom view.",
        ),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        _bold(p, name)
        p.add_run(f": {desc}")

    _heading(doc, "VIII-B. Four-Panel Pedagogy Viewer", 2)
    doc.add_paragraph("The SingleShuntDialog provides a four-panel animated view:")
    panels = [
        (
            "Slow-time modulation panel",
            "Shows the three-phase duty cycle envelopes (D_A, D_B, D_C) over one full electrical cycle "
            "with SVM sector colour bands and a time cursor.",
        ),
        (
            "SVM hexagon panel",
            "Displays the αβ-plane hexagon with the active sector highlighted and the reference "
            "vector rotating synchronously with the cursor position.",
        ),
        (
            "PWM period zoom panel",
            "Shows the selected period's three pulse waveforms at high resolution, with dead-time gaps "
            "highlighted in red and W₁, W₂ acquisition windows highlighted in green (observable) "
            "or red (blind).",
        ),
        (
            "Acquisition info panel",
            "Textual display: current sector, duty cycle ordering, W₁ and W₂ widths, "
            "which phase is measured, reconstruction feasibility, and KCL identity used.",
        ),
    ]
    for name, desc in panels:
        p = doc.add_paragraph(style="List Number")
        _bold(p, f"{name}: ")
        p.add_run(desc)

    buf9 = _fig_to_stream(fig_reconstruction_algorithm())
    doc.add_picture(buf9, width=Inches(4.5))
    _add_caption(
        doc,
        "Fig. 9 – Single Shunt Current Reconstruction Algorithm Flowchart "
        "implemented in svm_shaper/single_shunt.py.",
    )

    _heading(doc, "VIII-C. Compensation Strategy Selector", 2)
    doc.add_paragraph(
        "A drop-down menu in the SingleShuntDialog allows selecting the active compensation mode:"
    )
    comps = [
        "None (display blind zones only — pedagogic mode)",
        "Minimum pulse insertion (adaptive δ per period)",
        "Duty-cycle redistribution (symmetric δ/2 adjustment)",
        "Hold-last-value predictor (for deep blind zones)",
    ]
    for c in comps:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(c)
    doc.add_paragraph(
        "The selected strategy's duty-cycle perturbation is applied only to the time-domain "
        "visualisation and acquisition info within the dialog; it does not modify the main "
        "simulation result stored in SimulationResult."
    )

    # ---------- Section IX: Conclusion ----------
    _heading(doc, "IX. Conclusion", 1)
    doc.add_paragraph(
        "This report has presented a rigorous treatment of single shunt DC-bus current "
        "reconstruction for three-phase PWM inverters. The key contributions are:"
    )
    conclusions = [
        "A unified window-width formula (5)–(8) valid for all duty-cycle orderings, "
        "all modulation modes, and both center-aligned and edge-aligned carriers.",
        "A per-sector phase assignment table (Table III) mapping shunt measurements to "
        "physical phase currents for all six SVM sectors.",
        "A quantitative blind-zone analysis (13) estimating the angular extent of "
        "reconstruction failure as a function of modulation index, PWM frequency, and "
        "dead time.",
        "A comparative evaluation of SPWM, THIPWM, SVM, and six DPWM variants with "
        "respect to their inherent SSCR window-width profiles (Table IV, Fig. 7).",
        "Three compensation strategies for center-aligned PWM blind zones, with a "
        "voltage-error bound (14).",
        "A phase-shift strategy for edge-aligned PWM that guarantees observable windows "
        "independent of the duty cycle values (20).",
        "An integration plan for the SVM Analyst simulation tool, implementing the above "
        "as a dedicated pedagogical viewer (svm_shaper/single_shunt.py, SingleShuntDialog).",
    ]
    for c in conclusions:
        p = doc.add_paragraph(style="List Number")
        p.add_run(c)

    # ---------- Symbol Glossary ----------
    _heading(doc, "X. Symbol Glossary", 1)
    buf10 = _fig_to_stream(fig_symbol_table_figure(), dpi=160)
    doc.add_picture(buf10, width=Inches(6.0))
    _add_caption(doc, "Table I – List of Symbols and Notation Used in This Report.")

    # ---------- References ----------
    _heading(doc, "References", 1)
    refs = [
        "[1] P. Vas, Sensorless Vector and Direct Torque Control, Oxford University Press, 1998.",
        "[2] J.-S. Kim and S.-K. Sul, 'New approaches for high-performance pulsed-width "
        "modulation with one current sensor,' IEEE Trans. Power Electron., vol. 11, no. 3, "
        "pp. 409–417, May 1996.",
        "[3] H. W. van der Broeck, H.-C. Skudelny, and G. V. Stanke, 'Analysis and realization "
        "of a pulsewidth modulator based on voltage space vectors,' IEEE Trans. Ind. Appl., "
        "vol. 24, no. 1, pp. 142–150, Jan./Feb. 1988.",
        "[4] M. Corley and R. Lorenz, 'Rotor position and velocity estimation for a salient-pole "
        "permanent magnet synchronous machine at standstill and high speeds,' IEEE Trans. Ind. "
        "Appl., vol. 34, no. 4, pp. 784–789, Jul./Aug. 1998.",
        "[5] J. Holtz, 'Pulsewidth modulation – a survey,' IEEE Trans. Ind. Electron., vol. 39, "
        "no. 5, pp. 410–420, Oct. 1992.",
        "[6] H. Kim, T. Jahns, 'Phase current reconstruction for AC motor drives using a DC link "
        "single current sensor and measurement voltage vectors,' IEEE Trans. Power Electron., "
        "vol. 21, no. 5, pp. 1413–1419, Sep. 2006.",
        "[7] L. Woo-Cheol, L. Taek-Kie, and H. Dong-Seok, 'A three-phase current-mismatch "
        "problem in a 3-phase AC power conditioning system using a single current sensor,' "
        "IEEE Trans. Power Electron., vol. 15, no. 6, pp. 1049–1056, Nov. 2000.",
        "[8] D. G. Holmes and T. A. Lipo, Pulse Width Modulation for Power Converters: "
        "Principles and Practice. Wiley-IEEE Press, 2003.",
        "[9] A. Hava, R. Kerkman, and T. Lipo, 'Simple analytical and graphical methods for "
        "carrier-based PWM-VSI drives,' IEEE Trans. Power Electron., vol. 14, no. 1, "
        "pp. 49–61, Jan. 1999.",
        "[10] R. Teodorescu, M. Liserre, and P. Rodriguez, Grid Converters for Photovoltaic "
        "and Wind Power Systems. Wiley-IEEE Press, 2011.",
    ]
    for ref in refs:
        p = doc.add_paragraph(style="List Number")
        p.add_run(ref).font.size = Pt(9)

    doc.save(out_path)
    print(f"[OK] Report written → {out_path}")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    out = os.path.join(root, "docs", "SSCR_Technical_Report.docx")
    os.makedirs(os.path.dirname(out), exist_ok=True)
    build_document(out)
